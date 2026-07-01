import importlib.machinery
import sys
import tempfile
import types
import unittest
from io import BytesIO
from pathlib import Path
from unittest.mock import patch

from docx import Document

from contract_agent.config.config_parser import ParserConfig
from contract_agent.parser import (
    BlockLocation,
    ClauseChunk,
    ContractParser,
    DocumentBlock,
    DocumentLoadError,
    DocumentParseError,
    DocumentMetadata,
    DocumentSpan,
    DocumentTable,
    ParsedDocument,
    UnsupportedFileType,
    to_rag_documents,
)
from contract_agent.parser.markdown_document import MarkdownDocument
from contract_agent.parser.parsed.markdown_metadata_builder import build_metadata


class ContractParserServiceTests(unittest.TestCase):
    def test_metadata_doc_id_uses_raw_text_and_unknown_page_count_is_zero(self):
        first = build_metadata(
            file_name="same.txt",
            source_path="same.txt",
            file_type="txt",
            raw_text="first body",
            spans=[
                DocumentSpan(
                    span_id="s1",
                    block_index=0,
                    start_offset=0,
                    end_offset=5,
                    text="first",
                )
            ],
        )
        second = build_metadata(
            file_name="same.txt",
            source_path="same.txt",
            file_type="txt",
            raw_text="second body",
            spans=[
                DocumentSpan(
                    span_id="s1",
                    block_index=0,
                    start_offset=0,
                    end_offset=6,
                    text="second",
                )
            ],
        )

        self.assertNotEqual(first.doc_id, second.doc_id)
        self.assertEqual(first.page_count, 0)

    def test_parse_text_generates_metadata_spans_and_chunks_without_detectors(self):
        document = _builtin_parser().parse_text("第一条 付款\n甲方应支付价款。", "inline.txt")

        self.assertEqual(document.metadata.file_name, "inline.txt")
        self.assertEqual(document.metadata.file_type, "txt")
        self.assertEqual(document.raw_text, "第一条 付款 甲方应支付价款。")
        self.assertEqual(document.schema_version, "2.0")
        self.assertEqual(len(document.spans), 1)
        self.assertEqual(len(document.blocks), len(document.spans))
        self.assertFalse(hasattr(document, "detector_results"))
        self.assertIn("parser_backend", document.conversion_metadata)
        self.assertIn("第一条 付款", document.markdown_content)
        self.assertTrue(document.clause_chunks)
        self.assertTrue(
            any("甲方应支付价款" in chunk.source_text for chunk in document.clause_chunks)
        )
        self.assertTrue(
            any("甲方应支付价款" in item["page_content"] for item in to_rag_documents(document))
        )

    def test_parse_markdown_recognizes_common_list_markers(self):
        markdown = "\n".join(
            [
                "- dash item",
                "* star item",
                "+ plus item",
                "1. ordered item",
                "2) parenthesized item",
                "  - indented dash item",
                "   3. indented ordered item",
            ]
        )

        document = ContractParser().parse_markdown(
            MarkdownDocument(
                markdown_content=markdown,
                file_name="contract.md",
                file_type="md",
                source_path="inline",
                backend_name="builtin",
            )
        )

        self.assertEqual([block.block_type for block in document.blocks], ["list_item"] * 7)
        self.assertEqual(
            [block.text for block in document.blocks],
            [
                "dash item",
                "star item",
                "plus item",
                "ordered item",
                "parenthesized item",
                "indented dash item",
                "indented ordered item",
            ],
        )

        ambiguous_markdown = "\n".join(
            [
                "\\- escaped marker",
                "1.23 decimal value",
                "| - | table cell |",
            ]
        )

        ambiguous_document = ContractParser().parse_markdown(
            MarkdownDocument(
                markdown_content=ambiguous_markdown,
                file_name="contract.md",
                file_type="md",
                source_path="inline",
                backend_name="builtin",
            )
        )

        self.assertNotIn(
            "list_item",
            [block.block_type for block in ambiguous_document.blocks],
        )
        extra_ambiguous_document = ContractParser().parse_markdown(
            MarkdownDocument(
                markdown_content="    - code item\n---",
                file_name="contract.md",
                file_type="md",
                source_path="inline",
                backend_name="builtin",
            )
        )
        self.assertNotIn(
            "list_item",
            [block.block_type for block in extra_ambiguous_document.blocks],
        )

    def test_parse_bytes_supports_txt_encodings(self):
        parser = _builtin_parser()

        utf8 = parser.parse_bytes("a.txt", "第一条 付款".encode("utf-8"))
        utf8_sig = parser.parse_bytes("b.txt", "第一条 付款".encode("utf-8-sig"))
        gb18030 = parser.parse_bytes("c.txt", "第一条 付款".encode("gb18030"))
        utf16 = parser.parse_bytes("d.txt", "采购合同\n甲方：A公司".encode("utf-16"))

        self.assertEqual(utf8.raw_text, "第一条 付款")
        self.assertEqual(utf8_sig.raw_text, "第一条 付款")
        self.assertEqual(gb18030.raw_text, "第一条 付款")
        self.assertEqual(utf16.raw_text, "采购合同 甲方：A公司")

    def test_utf8_bom_does_not_pollute_first_txt_chunk(self):
        document = _builtin_parser().parse_bytes(
            "bom.txt",
            "\ufeff第一条 标的\n甲方：A\n乙方：B".encode("utf-8"),
        )

        self.assertNotIn("\ufeff", document.raw_text)
        self.assertEqual(document.clause_chunks[0].source_text, "第一条 标的 甲方：A 乙方：B")

    def test_corrupt_docx_and_pdf_raise_user_input_parser_errors(self):
        parser = _builtin_parser()

        with self.assertRaisesRegex(DocumentLoadError, "docx"):
            parser.parse_bytes("bad.docx", b"not a real file")
        with self.assertRaisesRegex(DocumentLoadError, "pdf"):
            parser.parse_bytes("bad.pdf", b"not a real file")

    def test_parse_docx_preserves_table_content_in_dom_markdown_chunks_and_rag(self):
        docx = Document()
        docx.add_paragraph("Project Application")
        table = docx.add_table(rows=3, cols=2)
        table.cell(0, 0).text = "Project Name"
        table.cell(0, 1).text = "Contract review with MCP and RAG"
        table.cell(1, 0).text = "Project Type"
        table.cell(1, 1).text = "Software"
        table.cell(2, 0).text = "Reason"
        table.cell(2, 1).text = "Improve contract review quality"
        buffer = BytesIO()
        docx.save(buffer)

        document = _builtin_parser().parse_bytes("project.docx", buffer.getvalue())

        self.assertIn("Project Name", document.raw_text)
        self.assertIn("Contract review with MCP and RAG", document.raw_text)
        self.assertEqual(len(document.tables), 1)
        self.assertEqual(
            document.tables[0].rows[0],
            ["Project Name", "Contract review with MCP and RAG"],
        )
        table_blocks = [block for block in document.blocks if block.block_type == "table"]
        self.assertEqual(len(table_blocks), 1)
        self.assertIn("| Project Name |", table_blocks[0].markdown or "")
        self.assertIn("Project Type | Software", document.markdown_content)
        self.assertTrue(any("Reason" in chunk.source_text for chunk in document.clause_chunks))
        self.assertTrue(
            all(
                document.raw_text[chunk.start_offset : chunk.end_offset] == chunk.source_text
                for chunk in document.clause_chunks
                if chunk.chunk_level != "table"
            )
        )
        self.assertTrue(
            any("Project Name" in item["page_content"] for item in to_rag_documents(document))
        )
        self.assertIsNotNone(document.semantic_graph)
        node_types = {node["type"] for node in document.semantic_graph.nodes}
        self.assertIn("document", node_types)
        self.assertIn("block", node_types)
        self.assertIn("chunk", node_types)
        self.assertIn("table", node_types)
        edge_types = {edge["type"] for edge in document.semantic_graph.edges}
        self.assertIn("contains", edge_types)
        self.assertIn("derived_from", edge_types)

    def test_parse_docx_preserves_adjacent_duplicate_table_cell_values(self):
        docx = Document()
        docx.add_paragraph("Penalty Table")
        table = docx.add_table(rows=1, cols=3)
        table.cell(0, 0).text = "Penalty"
        table.cell(0, 1).text = "10%"
        table.cell(0, 2).text = "10%"
        buffer = BytesIO()
        docx.save(buffer)

        document = _builtin_parser().parse_bytes("duplicate-cells.docx", buffer.getvalue())

        self.assertEqual(document.tables[0].rows[0], ["Penalty", "10%", "10%"])
        table_block = next(block for block in document.blocks if block.block_type == "table")
        self.assertIn("| Penalty | 10% | 10% |", table_block.markdown or "")

    def test_parse_docx_collapses_merged_table_cell_projection_by_xml_identity(self):
        docx = Document()
        docx.add_paragraph("Merged Table")
        table = docx.add_table(rows=1, cols=3)
        merged = table.cell(0, 0).merge(table.cell(0, 1))
        merged.text = "Merged label"
        table.cell(0, 2).text = "Value"
        buffer = BytesIO()
        docx.save(buffer)

        document = _builtin_parser().parse_bytes("merged-cells.docx", buffer.getvalue())

        self.assertEqual(document.tables[0].rows[0], ["Merged label", "", "Value"])
        table_block = next(block for block in document.blocks if block.block_type == "table")
        self.assertIn("| Merged label |  | Value |", table_block.markdown or "")

    def test_parse_docx_collapses_vertical_merged_cell_projection_by_xml_identity(self):
        docx = Document()
        docx.add_paragraph("Vertical Merge Table")
        table = docx.add_table(rows=2, cols=2)
        merged = table.cell(0, 0).merge(table.cell(1, 0))
        merged.text = "Vertical label"
        table.cell(0, 1).text = "Top"
        table.cell(1, 1).text = "Bottom"
        buffer = BytesIO()
        docx.save(buffer)

        document = _builtin_parser().parse_bytes("vertical-merged-cells.docx", buffer.getvalue())

        self.assertEqual(document.tables[0].rows, [["Vertical label", "Top"], ["", "Bottom"]])
        table_block = next(block for block in document.blocks if block.block_type == "table")
        self.assertIn("| Vertical label | Top |", table_block.markdown or "")
        self.assertIn("|  | Bottom |", table_block.markdown or "")

    def test_parse_docx_preserves_placeholders_for_rectangular_merged_cells(self):
        docx = Document()
        docx.add_paragraph("Rectangular Merge Table")
        table = docx.add_table(rows=2, cols=3)
        merged = table.cell(0, 0).merge(table.cell(1, 1))
        merged.text = "Merged block"
        table.cell(0, 2).text = "Top tail"
        table.cell(1, 2).text = "Bottom tail"
        buffer = BytesIO()
        docx.save(buffer)

        document = _builtin_parser().parse_bytes("rectangular-merged-cells.docx", buffer.getvalue())

        self.assertEqual(
            document.tables[0].rows,
            [["Merged block", "", "Top tail"], ["", "", "Bottom tail"]],
        )
        table_block = next(block for block in document.blocks if block.block_type == "table")
        self.assertIn("| Merged block |  | Top tail |", table_block.markdown or "")
        self.assertIn("|  |  | Bottom tail |", table_block.markdown or "")

    def test_parse_docx_keeps_all_placeholder_rows_for_fully_merged_rectangles(self):
        docx = Document()
        docx.add_paragraph("Fully Merged Rectangle")
        table = docx.add_table(rows=2, cols=2)
        merged = table.cell(0, 0).merge(table.cell(1, 1))
        merged.text = "Merged all"
        buffer = BytesIO()
        docx.save(buffer)

        document = _builtin_parser().parse_bytes("fully-merged-rectangle.docx", buffer.getvalue())

        self.assertEqual(document.tables[0].rows, [["Merged all", ""], ["", ""]])
        table_block = next(block for block in document.blocks if block.block_type == "table")
        self.assertIn("| Merged all |  |", table_block.markdown or "")
        self.assertIn("|  |  |", table_block.markdown or "")

    def test_parse_docx_rejects_unreasonable_grid_span(self):
        class FakeGridSpan:
            val = "101"

        class FakeCellProperties:
            gridSpan = FakeGridSpan()
            vMerge = None

        class FakeCellXml:
            tcPr = FakeCellProperties()

        class FakeRowXml:
            tc_lst = [FakeCellXml()]

        class FakeRow:
            _tr = FakeRowXml()

        class FakeTable:
            rows = [FakeRow()]

        from contract_agent.parser.convertor import builtin_markdown_converter as converter

        with self.assertRaises(DocumentParseError):
            converter._table_rows(FakeTable())

    def test_parse_docx_keeps_markdown_when_optional_html_conversion_fails(self):
        from contract_agent.parser.convertor import builtin_markdown_converter as converter

        with patch.object(converter, "_parse_docx_bytes", return_value="Docx markdown"):
            with patch.object(converter, "_docx_to_html", side_effect=RuntimeError("html failed")):
                loaded = converter.load_bytes("contract.docx", b"fake")

        self.assertEqual(loaded.markdown_content, "Docx markdown")
        self.assertEqual(loaded.html_content, "")

    def test_semantic_graph_chunk_edges_do_not_grow_by_blocks_times_chunks(self):
        spans = []
        blocks = []
        chunks = []
        cursor = 0
        for index in range(20):
            text = f"Block {index}"
            start = cursor
            end = start + len(text)
            span_id = f"p1-b{index}"
            spans.append(
                DocumentSpan(
                    span_id=span_id,
                    page_no=1,
                    block_index=index,
                    start_offset=start,
                    end_offset=end,
                    text=text,
                )
            )
            blocks.append(
                DocumentBlock(
                    block_id=span_id,
                    block_type="paragraph",
                    text=text,
                    location=BlockLocation(
                        page_no=1,
                        block_index=index,
                        start_offset=start,
                        end_offset=end,
                        span_ids=[span_id],
                    ),
                )
            )
            chunks.append(
                ClauseChunk(
                    chunk_id=f"chunk-{index}",
                    chunk_level="sentence_group",
                    section_title="Body",
                    page_no=1,
                    start_offset=0,
                    end_offset=2000,
                    source_text=text,
                )
            )
            cursor = end + 1
        document = ParsedDocument(
            metadata=DocumentMetadata(
                doc_id="doc-graph",
                file_name="graph.txt",
                file_type="txt",
                source_path="inline",
            ),
            raw_text="\n".join(span.text for span in spans),
            spans=spans,
            blocks=list(reversed(blocks)),
            clause_chunks=chunks,
        )

        graph = ContractParser()._build_semantic_graph(document)

        derived_edges = [edge for edge in graph.edges if edge["type"] == "derived_from"]
        self.assertLessEqual(len(derived_edges), len(chunks))

    def test_semantic_graph_exactly_links_chunks_with_matching_block_offsets(self):
        spans = []
        blocks = []
        chunks = []
        cursor = 0
        for index in range(20):
            text = f"Block {index}"
            start = cursor
            end = start + len(text)
            span_id = f"p1-b{index}"
            spans.append(
                DocumentSpan(
                    span_id=span_id,
                    page_no=1,
                    block_index=index,
                    start_offset=start,
                    end_offset=end,
                    text=text,
                )
            )
            blocks.append(
                DocumentBlock(
                    block_id=span_id,
                    block_type="paragraph",
                    text=text,
                    location=BlockLocation(
                        page_no=1,
                        block_index=index,
                        start_offset=start,
                        end_offset=end,
                        span_ids=[span_id],
                    ),
                )
            )
            chunks.append(
                ClauseChunk(
                    chunk_id=f"chunk-{index}",
                    chunk_level="sentence_group",
                    section_title="Body",
                    page_no=1,
                    start_offset=start,
                    end_offset=end,
                    source_text=text,
                )
            )
            cursor = end + 1
        document = ParsedDocument(
            metadata=DocumentMetadata(
                doc_id="doc-graph-exact",
                file_name="graph.txt",
                file_type="txt",
                source_path="inline",
            ),
            raw_text="\n".join(span.text for span in spans),
            spans=spans,
            blocks=list(reversed(blocks)),
            clause_chunks=chunks,
        )

        graph = ContractParser()._build_semantic_graph(document)

        derived_edges = [edge for edge in graph.edges if edge["type"] == "derived_from"]
        self.assertEqual(
            derived_edges,
            [
                {
                    "source": f"chunk:chunk-{index}",
                    "target": f"block:p1-b{index}",
                    "type": "derived_from",
                }
                for index in range(20)
            ],
        )

    def test_semantic_graph_links_chunks_by_offsets_independent_of_chunk_order(self):
        blocks = [
            DocumentBlock(
                block_id="block-a",
                block_type="paragraph",
                text="First block",
                location=BlockLocation(
                    page_no=1,
                    block_index=0,
                    start_offset=0,
                    end_offset=11,
                    span_ids=["span-a"],
                ),
            ),
            DocumentBlock(
                block_id="block-b",
                block_type="paragraph",
                text="Second block",
                location=BlockLocation(
                    page_no=1,
                    block_index=1,
                    start_offset=12,
                    end_offset=24,
                    span_ids=["span-b"],
                ),
            ),
        ]
        chunks = [
            ClauseChunk(
                chunk_id="chunk-b",
                chunk_level="paragraph",
                section_title="Second block",
                page_no=1,
                start_offset=12,
                end_offset=24,
                source_text="Second block",
            ),
            ClauseChunk(
                chunk_id="chunk-a",
                chunk_level="paragraph",
                section_title="First block",
                page_no=1,
                start_offset=0,
                end_offset=11,
                source_text="First block",
            ),
        ]
        document = ParsedDocument(
            metadata=DocumentMetadata(
                doc_id="doc-graph-order",
                file_name="graph.txt",
                file_type="txt",
                source_path="inline",
            ),
            raw_text="First block\nSecond block",
            blocks=blocks,
            clause_chunks=chunks,
        )

        graph = ContractParser()._build_semantic_graph(document)

        derived_targets = {
            (edge["source"], edge["target"])
            for edge in graph.edges
            if edge["type"] == "derived_from" and str(edge["source"]).startswith("chunk:")
        }
        self.assertEqual(
            derived_targets,
            {
                ("chunk:chunk-a", "block:block-a"),
                ("chunk:chunk-b", "block:block-b"),
            },
        )

    def test_semantic_graph_preserves_block_table_and_chunk_page_numbers(self):
        document = ParsedDocument(
            metadata=DocumentMetadata(
                doc_id="doc-page-graph",
                file_name="graph.txt",
                file_type="txt",
                source_path="inline",
            ),
            raw_text="Unknown\nKnown",
            blocks=[
                DocumentBlock(
                    block_id="block-unknown",
                    block_type="paragraph",
                    text="Unknown",
                    location=BlockLocation(
                        page_no=None,
                        block_index=0,
                        start_offset=0,
                        end_offset=7,
                        span_ids=["span-unknown"],
                    ),
                ),
                DocumentBlock(
                    block_id="block-known",
                    block_type="table",
                    text="Known",
                    location=BlockLocation(
                        page_no=2,
                        block_index=1,
                        start_offset=8,
                        end_offset=13,
                        span_ids=["span-known"],
                    ),
                ),
            ],
            tables=[DocumentTable(table_id="table-known", page_no=2, span_ids=["span-known"])],
            clause_chunks=[
                ClauseChunk(
                    chunk_id="chunk-unknown",
                    chunk_level="paragraph",
                    section_title="Unknown",
                    page_no=None,
                    start_offset=0,
                    end_offset=7,
                    source_text="Unknown",
                ),
                ClauseChunk(
                    chunk_id="chunk-known",
                    chunk_level="table",
                    section_title="Known",
                    page_no=2,
                    start_offset=8,
                    end_offset=13,
                    source_text="Known",
                ),
            ],
        )

        graph = ContractParser()._build_semantic_graph(document)
        page_by_node = {node["id"]: node["metadata"].get("page_no") for node in graph.nodes}

        self.assertIsNone(page_by_node["block:block-unknown"])
        self.assertEqual(page_by_node["block:block-known"], 2)
        self.assertEqual(page_by_node["table:table-known"], 2)
        self.assertIsNone(page_by_node["chunk:chunk-unknown"])
        self.assertEqual(page_by_node["chunk:chunk-known"], 2)

    def test_semantic_graph_definition_node_ids_are_unique_for_repeated_terms(self):
        from contract_agent.parser import DocumentDefinition

        document = ParsedDocument(
            metadata=DocumentMetadata(
                doc_id="doc-definitions",
                file_name="definitions.txt",
                file_type="txt",
                source_path="inline",
            ),
            raw_text="Definitions",
            definitions=[
                DocumentDefinition(term="Affiliate", definition="First meaning"),
                DocumentDefinition(term="Affiliate", definition="Second meaning"),
            ],
        )

        graph = ContractParser()._build_semantic_graph(document)

        definition_ids = [node["id"] for node in graph.nodes if node.get("type") == "definition"]
        self.assertEqual(definition_ids, ["definition:0:Affiliate", "definition:1:Affiliate"])
        self.assertEqual(len(definition_ids), len(set(definition_ids)))

    def test_parse_path_delegates_to_bytes_loader(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "contract.txt"
            path.write_text("第一条 付款", encoding="utf-8")

            parser = ContractParser(
                parser_config=ParserConfig(
                    default_converter="builtin",
                    enabled_converters=["builtin"],
                    fallback_order=["builtin"],
                    allow_path_input=True,
                    trusted_path_roots=[tmp],
                )
            )
            document = parser.parse_path(path)

        self.assertEqual(document.metadata.file_name, "contract.txt")
        self.assertEqual(document.raw_text, "第一条 付款")

    def test_convert_to_markdown_exposes_exact_docling_markdown_before_parsing(self):
        markdown = "# Contract\n\n| Key | Value |\n| --- | --- |\n"

        class FakeDoclingDocument:
            def export_to_markdown(self, **kwargs: object) -> str:
                return markdown

        class FakeBackendResult:
            document = FakeDoclingDocument()

        class FakeDocumentConverter:
            def __init__(self, **kwargs: object) -> None:
                self.kwargs = kwargs

            def convert(self, source: str) -> FakeBackendResult:
                return FakeBackendResult()

        with tempfile.TemporaryDirectory() as tmp:
            config = ParserConfig(
                default_converter="docling",
                enabled_converters=["docling"],
                fallback_order=["docling"],
                docling_enabled=True,
                allow_path_input=True,
                trusted_path_roots=[tmp],
            )
            path = Path(tmp) / "contract.pdf"
            path.write_bytes(b"fake")
            with (
                _fake_docling_modules(FakeDocumentConverter),
                patch("importlib.util.find_spec", return_value=_module_spec("docling")),
            ):
                parser = ContractParser(parser_config=config)
                markdown_document = parser.convert_to_markdown(path)
                parsed_document = parser.parse_markdown(markdown_document)

        self.assertEqual(markdown_document.markdown_content, markdown)
        self.assertEqual(markdown_document.backend_name, "docling")
        self.assertEqual(parsed_document.markdown_content, markdown)

    def test_parse_path_routes_file_through_docling_backend_and_outputs_chunks_for_rag(self):
        class FakeCell:
            def __init__(self, text: str) -> None:
                self.text = text

        class FakeTableData:
            grid = [[FakeCell("Key"), FakeCell("Value")], [FakeCell("Project"), FakeCell("Alpha")]]

        class FakeTable:
            self_ref = "#/tables/0"
            data = FakeTableData()
            prov = []

            def caption_text(self, document: object) -> str:
                return "Project table"

            def export_to_markdown(self, document: object) -> str:
                return "| Key | Value |\n| --- | --- |\n| Project | Alpha |"

        class FakeDoclingDocument:
            tables = [FakeTable()]

            def export_to_markdown(self, **kwargs: object) -> str:
                return "# Contract\n\n| Key | Value |\n| --- | --- |\n| Project | Alpha |"

        class FakeBackendResult:
            document = FakeDoclingDocument()

        class FakeDocumentConverter:
            def __init__(self, **kwargs: object) -> None:
                self.kwargs = kwargs

            def convert(self, source: str) -> FakeBackendResult:
                return FakeBackendResult()

        with tempfile.TemporaryDirectory() as tmp:
            config = ParserConfig(
                default_converter="docling",
                enabled_converters=["docling"],
                fallback_order=["docling"],
                docling_enabled=True,
                allow_path_input=True,
                trusted_path_roots=[tmp],
            )
            path = Path(tmp) / "contract.pdf"
            path.write_bytes(b"fake")
            with (
                _fake_docling_modules(FakeDocumentConverter),
                patch("importlib.util.find_spec", return_value=_module_spec("docling")),
            ):
                document = ContractParser(parser_config=config).parse_path(path)

        self.assertEqual(document.conversion_metadata["parser_backend"], "docling")
        self.assertEqual(document.tables[0].rows, [["Key", "Value"], ["Project", "Alpha"]])
        self.assertTrue(document.clause_chunks)
        self.assertTrue(any("Alpha" in chunk.source_text for chunk in document.clause_chunks))
        self.assertTrue(any("Alpha" in item["page_content"] for item in to_rag_documents(document)))

    def test_parse_path_routes_legacy_doc_file_through_markitdown_backend(self):
        markdown = "# Legacy Contract\n\n| Key | Value |\n| --- | --- |\n| Project | Beta |\n"

        class FakeResult:
            text_content = markdown

        class FakeMarkItDown:
            def convert(self, source: str) -> FakeResult:
                self.source = source
                return FakeResult()

        module = types.ModuleType("markitdown")
        module.MarkItDown = FakeMarkItDown
        with tempfile.TemporaryDirectory() as tmp:
            config = ParserConfig(
                default_converter="markitdown",
                enabled_converters=["markitdown"],
                fallback_order=["markitdown"],
                markitdown_enabled=True,
                allowed_suffixes=[".doc"],
                allow_path_input=True,
                trusted_path_roots=[tmp],
            )
            path = Path(tmp) / "legacy.doc"
            path.write_bytes(b"fake legacy doc")
            with (
                patch.dict(sys.modules, {"markitdown": module}),
                patch("importlib.util.find_spec", return_value=_module_spec("markitdown")),
            ):
                parser = ContractParser(parser_config=config)
                markdown_document = parser.convert_to_markdown(path)
                document = parser.parse_markdown(markdown_document)

        self.assertEqual(markdown_document.backend_name, "markitdown")
        self.assertEqual(markdown_document.markdown_content, markdown)
        self.assertEqual(document.markdown_content, markdown)
        self.assertEqual(document.tables[0].rows, [["Key", "Value"], ["Project", "Beta"]])
        self.assertTrue(any("Beta" in item["page_content"] for item in to_rag_documents(document)))

    def test_parse_text_obeys_max_input_bytes(self):
        parser = ContractParser(parser_config=ParserConfig(max_input_bytes=8))

        with self.assertRaises(DocumentLoadError):
            parser.parse_text("123456789")

    def test_parse_markdown_obeys_max_input_bytes(self):
        parser = ContractParser(parser_config=ParserConfig(max_input_bytes=8))
        markdown_document = MarkdownDocument(
            markdown_content="123456789",
            file_name="inline.md",
            file_type="md",
            source_path="inline.md",
            backend_name="test",
        )

        with self.assertRaises(DocumentLoadError):
            parser.parse_markdown(markdown_document)

    def test_unsupported_suffix_raises_parser_exception(self):
        with self.assertRaises(UnsupportedFileType):
            _builtin_parser().parse_bytes("contract.xlsx", b"data")

    def test_empty_or_undecodable_text_raises_clear_parser_exception(self):
        parser = _builtin_parser()

        with self.assertRaises(DocumentParseError):
            parser.parse_text("   \n  ")
        with self.assertRaises(DocumentLoadError):
            parser.parse_bytes("contract.txt", b"\xff")


def _module_spec(name: str) -> importlib.machinery.ModuleSpec:
    return importlib.machinery.ModuleSpec(name, loader=None)


def _builtin_parser() -> ContractParser:
    return ContractParser(
        parser_config=ParserConfig(
            default_converter="builtin",
            enabled_converters=["builtin"],
            fallback_order=["builtin"],
        )
    )


def _fake_docling_modules(document_converter_cls: type):
    package = types.ModuleType("docling")
    module = types.ModuleType("docling.document_converter")
    base_models = types.ModuleType("docling.datamodel.base_models")
    pipeline_options = types.ModuleType("docling.datamodel.pipeline_options")

    class FakeInputFormat:
        PDF = "pdf"

    class RapidOcrOptions:
        def __init__(
            self,
            *,
            lang: list[str],
            force_full_page_ocr: bool,
            bitmap_area_threshold: float,
            text_score: float,
        ) -> None:
            self.lang = lang
            self.force_full_page_ocr = force_full_page_ocr
            self.bitmap_area_threshold = bitmap_area_threshold
            self.text_score = text_score

    class FakePdfPipelineOptions:
        def __init__(
            self,
            *,
            do_ocr: bool,
            ocr_options: RapidOcrOptions,
            do_table_structure: bool,
            ocr_batch_size: int,
            layout_batch_size: int,
            table_batch_size: int,
        ) -> None:
            self.do_ocr = do_ocr
            self.ocr_options = ocr_options
            self.do_table_structure = do_table_structure
            self.ocr_batch_size = ocr_batch_size
            self.layout_batch_size = layout_batch_size
            self.table_batch_size = table_batch_size

    class FakePdfFormatOption:
        def __init__(self, *, pipeline_options: FakePdfPipelineOptions) -> None:
            self.pipeline_options = pipeline_options

    module.DocumentConverter = document_converter_cls
    module.PdfFormatOption = FakePdfFormatOption
    base_models.InputFormat = FakeInputFormat
    pipeline_options.PdfPipelineOptions = FakePdfPipelineOptions
    pipeline_options.RapidOcrOptions = RapidOcrOptions
    return patch.dict(
        sys.modules,
        {
            "docling": package,
            "docling.document_converter": module,
            "docling.datamodel.base_models": base_models,
            "docling.datamodel.pipeline_options": pipeline_options,
        },
    )


if __name__ == "__main__":
    unittest.main()
