import importlib.util
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document

from contract_agent.config.config_parser import ParserConfig
from contract_agent.parser import ContractParser


class DoclingSmokeTests(unittest.TestCase):
    def test_txt_routes_to_builtin_by_default(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            path = root / "contract.txt"
            path.write_text("第一条 付款", encoding="utf-8")
            parser = ContractParser(parser_config=_smoke_config(root))

            markdown = parser.convert_path(path)

        self.assertEqual(markdown.backend_name, "builtin")
        self.assertIn("付款", markdown.markdown_content)

    @unittest.skipIf(importlib.util.find_spec("docling") is None, "docling is not installed")
    def test_docling_converts_real_structured_text_samples(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            samples = _write_structured_samples(root)
            parser = ContractParser(parser_config=_smoke_config(root))

            for suffix, path in samples.items():
                with self.subTest(suffix=suffix):
                    markdown = parser.convert_path(path)

                    self.assertEqual(markdown.backend_name, "docling")
                    self.assertTrue(markdown.markdown_content.strip())
                    self.assertIn("Project", markdown.markdown_content)
                    self.assertIn("Alpha", markdown.markdown_content)
                    self.assertEqual(
                        markdown.conversion_metadata["docling_input_format"],
                        {
                            ".docx": "DOCX",
                            ".md": "MD",
                            ".html": "HTML",
                            ".csv": "CSV",
                            ".xlsx": "XLSX",
                        }[suffix],
                    )


def _smoke_config(root: Path) -> ParserConfig:
    return ParserConfig(
        allow_path_input=True,
        trusted_path_roots=[str(root)],
        docling_enable_ocr=False,
    )


def _write_structured_samples(root: Path) -> dict[str, Path]:
    docx_path = root / "contract.docx"
    doc = Document()
    doc.add_paragraph("Docx Contract")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Key"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Project"
    table.cell(1, 1).text = "Alpha"
    doc.save(docx_path)

    md_path = root / "contract.md"
    md_path.write_text(
        "# Markdown Contract\n\n| Key | Value |\n| --- | --- |\n| Project | Alpha |\n",
        encoding="utf-8",
    )

    html_path = root / "contract.html"
    html_path.write_text(
        "<html><body><h1>HTML Contract</h1><table><tr><td>Project</td><td>Alpha</td></tr></table></body></html>",
        encoding="utf-8",
    )

    csv_path = root / "contract.csv"
    csv_path.write_text("Key,Value\nProject,Alpha\n", encoding="utf-8")

    xlsx_path = root / "contract.xlsx"
    _write_minimal_xlsx(xlsx_path)

    return {
        ".docx": docx_path,
        ".md": md_path,
        ".html": html_path,
        ".csv": csv_path,
        ".xlsx": xlsx_path,
    }


def _write_minimal_xlsx(path: Path) -> None:
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr(
            "[Content_Types].xml",
            """<?xml version="1.0" encoding="UTF-8"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>
  <Override PartName="/xl/worksheets/sheet1.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>
</Types>""",
        )
        archive.writestr(
            "_rels/.rels",
            """<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/>
</Relationships>""",
        )
        archive.writestr(
            "xl/_rels/workbook.xml.rels",
            """<?xml version="1.0" encoding="UTF-8"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" Target="worksheets/sheet1.xml"/>
</Relationships>""",
        )
        archive.writestr(
            "xl/workbook.xml",
            """<?xml version="1.0" encoding="UTF-8"?>
<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
  <sheets>
    <sheet name="Sheet1" sheetId="1" r:id="rId1"/>
  </sheets>
</workbook>""",
        )
        archive.writestr(
            "xl/worksheets/sheet1.xml",
            """<?xml version="1.0" encoding="UTF-8"?>
<worksheet xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main">
  <sheetData>
    <row r="1">
      <c r="A1" t="inlineStr"><is><t>Key</t></is></c>
      <c r="B1" t="inlineStr"><is><t>Value</t></is></c>
    </row>
    <row r="2">
      <c r="A2" t="inlineStr"><is><t>Project</t></is></c>
      <c r="B2" t="inlineStr"><is><t>Alpha</t></is></c>
    </row>
  </sheetData>
</worksheet>""",
        )


if __name__ == "__main__":
    unittest.main()
