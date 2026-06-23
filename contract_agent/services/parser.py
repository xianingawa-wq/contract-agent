from __future__ import annotations

import hashlib
import re
from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader
from mammoth import convert_to_html as mammoth_to_html

from contract_agent.schemas.document import DocumentMetadata, DocumentSpan, ParsedDocument
from contract_agent.services.chunker import ContractChunker


class ContractParser:
    SUPPORTED_SUFFIXES = {".txt", ".docx", ".pdf"}

    def __init__(self) -> None:
        self.chunker = ContractChunker()

    def parse(self, file_path: str) -> ParsedDocument:
        path = Path(file_path).expanduser().resolve()
        suffix = path.suffix.lower()

        if suffix not in self.SUPPORTED_SUFFIXES:
            raise ValueError(f"Unsupported file type: {suffix}")
        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        return self.parse_bytes(path.name, path.read_bytes(), source_path=str(path))

    def parse_bytes(
        self, file_name: str, content: bytes, source_path: str | None = None
    ) -> ParsedDocument:
        suffix = Path(file_name).suffix.lower()
        if suffix not in self.SUPPORTED_SUFFIXES:
            raise ValueError(f"Unsupported file type: {suffix}")

        if suffix == ".txt":
            raw_text, spans = self._parse_txt_bytes(content)
        elif suffix == ".docx":
            raw_text, spans = self._parse_docx_bytes(content)
        else:
            raw_text, spans = self._parse_pdf_bytes(content)

        document = self._build_document(
            file_name, source_path or file_name, suffix, raw_text, spans
        )

        # Generate rich HTML for .docx files (headings, bold, tables, lists, etc.)
        if suffix == ".docx":
            document.html_content = self._docx_to_html(content)

        return document

    def parse_text(self, text: str, source_name: str = "inline.txt") -> ParsedDocument:
        raw_text, spans = self._build_spans_from_blocks(text.splitlines(), page_no=1)
        return self._build_document(source_name, source_name, ".txt", raw_text, spans)

    def _parse_txt_bytes(self, content: bytes) -> tuple[str, list[DocumentSpan]]:
        encodings = ("utf-8", "utf-8-sig", "gb18030")
        last_error: UnicodeDecodeError | None = None
        for encoding in encodings:
            try:
                text = content.decode(encoding)
                return self._build_spans_from_blocks(text.splitlines(), page_no=1)
            except UnicodeDecodeError as exc:
                last_error = exc
        raise ValueError("Unable to decode txt file with utf-8/utf-8-sig/gb18030.") from last_error

    def _parse_docx_bytes(self, content: bytes) -> tuple[str, list[DocumentSpan]]:
        doc = Document(BytesIO(content))
        blocks = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.text.strip()]
        return self._build_spans_from_blocks(blocks, page_no=1)

    def _parse_pdf_bytes(self, content: bytes) -> tuple[str, list[DocumentSpan]]:
        reader = PdfReader(BytesIO(content))
        all_spans: list[DocumentSpan] = []
        raw_parts: list[str] = []
        global_offset = 0

        for page_no, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if not text:
                continue

            blocks = [line.strip() for line in text.splitlines() if line.strip()]
            page_text, page_spans = self._build_spans_from_blocks(blocks, page_no=page_no)

            for span in page_spans:
                span.start_offset += global_offset
                span.end_offset += global_offset
                all_spans.append(span)

            raw_parts.append(page_text)
            global_offset += len(page_text) + 2

        return "\n\n".join(raw_parts), all_spans

    def _build_spans_from_blocks(
        self,
        blocks: list[str],
        page_no: int | None,
    ) -> tuple[str, list[DocumentSpan]]:
        spans: list[DocumentSpan] = []
        raw_parts: list[str] = []
        cursor = 0

        for block_index, block in enumerate(blocks):
            normalized = self._normalize_text(block)
            if not normalized:
                continue

            start_offset = cursor
            end_offset = start_offset + len(normalized)
            span_id = f"p{page_no or 0}-b{block_index}"

            spans.append(
                DocumentSpan(
                    span_id=span_id,
                    page_no=page_no,
                    block_index=block_index,
                    start_offset=start_offset,
                    end_offset=end_offset,
                    text=normalized,
                )
            )
            raw_parts.append(normalized)
            cursor = end_offset + 1

        return "\n".join(raw_parts), spans

    def _build_document(
        self,
        file_name: str,
        source_path: str,
        suffix: str,
        raw_text: str,
        spans: list[DocumentSpan],
    ) -> ParsedDocument:
        metadata = DocumentMetadata(
            doc_id=self._doc_id(source_path),
            file_name=file_name,
            file_type=suffix.lstrip("."),
            source_path=source_path,
            title=self._extract_title(raw_text, file_name),
            contract_type_hint=self._infer_contract_type(raw_text),
            party_a=self._extract_party(raw_text, "甲方"),
            party_b=self._extract_party(raw_text, "乙方"),
            signed_date=self._extract_signed_date(raw_text),
            page_count=max((span.page_no or 0) for span in spans) if spans else 1,
        )
        document = ParsedDocument(metadata=metadata, raw_text=raw_text, spans=spans)
        document.clause_chunks = self.chunker.chunk(document)
        return document

    def _normalize_text(self, text: str) -> str:
        return re.sub(r"\s+", " ", text).strip()

    def _extract_title(self, raw_text: str, fallback: str) -> str:
        for line in raw_text.splitlines():
            clean = line.strip()
            if clean:
                return clean[:100]
        return fallback

    def _extract_party(self, text: str, role: str) -> str | None:
        match = re.search(rf"{role}[:：]\s*(.+)", text)
        return match.group(1).strip() if match else None

    def _extract_signed_date(self, text: str) -> str | None:
        match = re.search(r"(\d{4}年\d{1,2}月\d{1,2}日)", text)
        return match.group(1) if match else None

    def _infer_contract_type(self, text: str) -> str | None:
        if "采购" in text:
            return "采购合同"
        if "保密" in text:
            return "保密协议"
        if "服务" in text:
            return "服务合同"
        return None

    def _docx_to_html(self, content: bytes) -> str:
        """Convert .docx to HTML using mammoth, preserving headings, bold, tables, lists, etc."""
        result = mammoth_to_html(BytesIO(content))
        if result.value:
            return result.value.strip()
        return ""

    def _doc_id(self, source: str) -> str:
        digest = hashlib.md5(source.encode("utf-8")).hexdigest()[:12]
        return f"doc_{digest}"
